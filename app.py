import streamlit as st

# MUST be the first Streamlit command
st.set_page_config(page_title="AI-Powered Application Assistant", page_icon="🤖", layout="wide")

import json
import os
from datetime import datetime
from anthropic import Anthropic
from dotenv import load_dotenv
import PyPDF2
from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF
import io
from supabase import create_client, Client

# Load environment variables
load_dotenv()

# Initialize Supabase client
@st.cache_resource
def init_supabase():
    """Initialize Supabase client with credentials from secrets or env."""
    url = st.secrets.get("SUPABASE_URL", os.getenv("SUPABASE_URL"))
    key = st.secrets.get("SUPABASE_KEY", os.getenv("SUPABASE_KEY"))
    return create_client(url, key)

supabase: Client = init_supabase()


# ===== AUTHENTICATION FUNCTIONS =====

def check_auth():
    """Check if user is authenticated. Returns user_id if authenticated, None otherwise."""
    if "user" not in st.session_state:
        return None
    return st.session_state["user"].id


def login_user(email, password):
    """Login user with email and password."""
    try:
        response = supabase.auth.sign_in_with_password({"email": email, "password": password})
        st.session_state["user"] = response.user
        return True, "Login successful!"
    except Exception as e:
        return False, str(e)


def signup_user(email, password):
    """Sign up new user with email and password."""
    try:
        response = supabase.auth.sign_up({"email": email, "password": password})
        if response.user:
            # Create profile entry
            supabase.table("profiles").insert({
                "id": response.user.id,
                "linkedin_url": "",
                "github_url": "",
                "portfolio_url": "",
                "candidate_name": "",
                "candidate_address": ""
            }).execute()
            return True, "Account created successfully! You can now log in."
        return False, "Failed to create account"
    except Exception as e:
        return False, str(e)


def logout_user():
    """Logout current user."""
    try:
        supabase.auth.sign_out()
        if "user" in st.session_state:
            del st.session_state["user"]
        return True
    except:
        return False


def show_auth_page():
    """Display login/signup page with guest option."""
    st.title("AI-Powered Application Assistant")
    st.markdown("### Welcome! Generate cover letters and answer application questions.")

    st.info("Create an account to save your resumes, cover letters, and track your applications. Or continue as guest to use the app without saving.")

    tab1, tab2, tab3 = st.tabs(["Continue as Guest", "Login", "Sign Up"])

    with tab1:
        st.subheader("Guest Mode")
        st.markdown("""
        **Try the app without creating an account:**
        - Generate cover letters
        - Answer application questions
        - Download your outputs

        **To save your work and access history, create an account!**
        """)

        if st.button("Continue as Guest", type="primary", use_container_width=True):
            st.session_state["guest_mode"] = True
            st.rerun()

    with tab2:
        st.subheader("Login")
        email = st.text_input("Email", key="login_email")
        password = st.text_input("Password", type="password", key="login_password")

        if st.button("Login", type="primary", use_container_width=True):
            if email and password:
                success, message = login_user(email, password)
                if success:
                    st.success(message)
                    if "guest_mode" in st.session_state:
                        del st.session_state["guest_mode"]
                    st.rerun()
                else:
                    st.error(f"Login failed: {message}")
            else:
                st.error("Please enter both email and password")

    with tab3:
        st.subheader("Create Account")
        new_email = st.text_input("Email", key="signup_email")
        new_password = st.text_input("Password", type="password", key="signup_password")
        confirm_password = st.text_input("Confirm Password", type="password", key="confirm_password")

        if st.button("Sign Up", type="primary", use_container_width=True):
            if new_email and new_password and confirm_password:
                if new_password != confirm_password:
                    st.error("Passwords do not match")
                elif len(new_password) < 6:
                    st.error("Password must be at least 6 characters")
                else:
                    success, message = signup_user(new_email, new_password)
                    if success:
                        st.success(message)
                    else:
                        st.error(f"Sign up failed: {message}")
            else:
                st.error("Please fill in all fields")


# Files for data storage (legacy - keeping for now)
RESUME_FILE = "resumes.json"
COVER_LETTER_FILE = "cover_letters.json"
RATINGS_FILE = "ratings.json"
PROFILE_FILE = "profile.json"
RESUME_FOLDER = "saved_resumes"

# Create resume folder if it doesn't exist
if not os.path.exists(RESUME_FOLDER):
    os.makedirs(RESUME_FOLDER)


# ===== DATABASE FUNCTIONS =====

def load_profile(user_id):
    """Load user profile from Supabase."""
    try:
        response = supabase.table("profiles").select("*").eq("id", user_id).execute()
        if response.data and len(response.data) > 0:
            return response.data[0]
        # Return default profile if none exists
        return {
            "linkedin_url": "",
            "github_url": "",
            "portfolio_url": "",
            "candidate_name": "",
            "candidate_address": "",
            "default_length": "concise",
            "default_tone": "conversational"
        }
    except:
        return {
            "linkedin_url": "",
            "github_url": "",
            "portfolio_url": "",
            "candidate_name": "",
            "candidate_address": "",
            "default_length": "concise",
            "default_tone": "conversational"
        }


def save_profile(user_id, profile_data):
    """Save user profile to Supabase."""
    try:
        # Check if profile exists
        existing = supabase.table("profiles").select("id").eq("id", user_id).execute()

        profile_data["id"] = user_id

        if existing.data and len(existing.data) > 0:
            # Update existing profile
            supabase.table("profiles").update(profile_data).eq("id", user_id).execute()
        else:
            # Insert new profile
            supabase.table("profiles").insert(profile_data).execute()
        return True
    except Exception as e:
        st.error(f"Error saving profile: {str(e)}")
        return False


def load_resumes(user_id):
    """Load all saved resumes from Supabase for this user."""
    try:
        response = supabase.table("resumes").select("*").eq("user_id", user_id).order("date_saved", desc=True).execute()
        return response.data if response.data else []
    except:
        return []


def save_resume(user_id, resume_data):
    """Save a new resume to Supabase."""
    try:
        resume_data["user_id"] = user_id
        supabase.table("resumes").insert(resume_data).execute()
        return True
    except Exception as e:
        st.error(f"Error saving resume: {str(e)}")
        return False


def load_cover_letters(user_id):
    """Load all saved cover letters from Supabase for this user."""
    try:
        response = supabase.table("cover_letters").select("*").eq("user_id", user_id).order("date_created", desc=True).execute()
        return response.data if response.data else []
    except:
        return []


def save_cover_letter(user_id, cover_letter_data):
    """Save a cover letter to Supabase."""
    try:
        cover_letter_data["user_id"] = user_id
        supabase.table("cover_letters").insert(cover_letter_data).execute()
        return True
    except Exception as e:
        st.error(f"Error saving cover letter: {str(e)}")
        return False


def delete_cover_letter(user_id, cover_letter_id):
    """Delete a cover letter from Supabase."""
    try:
        supabase.table("cover_letters").delete().eq("id", cover_letter_id).eq("user_id", user_id).execute()
        return True
    except Exception as e:
        st.error(f"Error deleting cover letter: {str(e)}")
        return False


def save_rating(user_id, rating_data):
    """Save a cover letter rating for ML training to Supabase."""
    try:
        rating_data["user_id"] = user_id
        supabase.table("ratings").insert(rating_data).execute()
        return True
    except Exception as e:
        st.error(f"Error saving rating: {str(e)}")
        return False


def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file."""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def extract_text_from_docx(docx_file):
    """Extract text from uploaded DOCX file."""
    doc = Document(docx_file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text


def get_latest_resume(user_id):
    """Get the most recently saved resume."""
    resumes = load_resumes(user_id)
    if resumes:
        return resumes[0]  # Already sorted by date_saved desc in load_resumes
    return None


def generate_cover_letter(resume_text, candidate_name, candidate_address, company_name, role_title, why_want_job, job_description="", additional_context="", resume_highlight="", length="concise", tone="conversational"):
    """Generate a cover letter using Claude Haiku API."""

    # Length instructions
    length_instructions = {
        "concise": "Keep the cover letter concise and focused, between 200-325 words. Be direct and impactful.",
        "standard": "Write a standard-length cover letter, between 325-450 words. Provide more detail while staying focused."
    }

    # Tone instructions
    tone_instructions = {
        "conversational": "Use a warm, conversational tone that is professional but approachable. Write as if speaking to a colleague. Avoid overly formal language while maintaining respect.",
        "professional": "Use a formal, traditional tone. Choose sophisticated vocabulary, avoid contractions, and maintain a serious, business-like demeanor throughout. This is for corporate, finance, law, or government roles.",
        "enthusiastic": "Use an energetic, passionate tone that shows genuine excitement about the role and company. Express enthusiasm naturally without going overboard. Perfect for startups, creative roles, or mission-driven organizations.",
        "confident": "Use a bold, direct tone that emphasizes your unique value proposition. Be assertive about your capabilities without arrogance. Focus on what you bring to the table. Ideal for competitive roles and leadership positions."
    }

    # Build prompt with XML tags for better structure and clarity
    prompt = f"""<instructions>
<length_requirement>
{length_instructions.get(length, length_instructions["concise"])}
</length_requirement>

<tone_requirement>
{tone_instructions.get(tone, tone_instructions["conversational"])}
</tone_requirement>

<additional_requirements>
- Do not use emojis
- Make the letter specific to this candidate and company
- Use concrete examples from the resume
- Do not include any XML tags, brackets, or meta-instructions in your output
- Output only the final cover letter text
</additional_requirements>
</instructions>

<resume>
{resume_text}
</resume>"""

    # Add resume highlight if provided
    if resume_highlight:
        prompt += f"""

<resume_highlight>
The candidate specifically wants to EMPHASIZE these experiences/achievements from their resume:

{resume_highlight}

IMPORTANT: Make sure to feature and highlight these specific items in the cover letter when relevant.
</resume_highlight>"""

    prompt += f"""

<job_description>
{job_description if job_description else "No job description provided. Focus on general fit with the company and role."}
</job_description>

<candidate_motivation>
{why_want_job}
</candidate_motivation>

<additional_context>
{additional_context if additional_context else "No additional context provided."}
</additional_context>

<output_format>
The cover letter must follow this exact structure:

{datetime.now().strftime("%B %d, %Y")}

{candidate_address}


Hiring Manager
{company_name}

Dear Hiring Manager,

[First paragraph: State why you are writing and include the exact title of the position: {role_title}. If applicable, mention any company connections.]

[Second paragraph: Describe what the candidate offers based on their resume. Provide specific examples of how their qualifications match the job requirements. Use work, classroom, or organizational experiences. Expand on resume details without repeating them verbatim.]

[Third paragraph: Establish synergy between the candidate and {company_name}. Include values, traits, corporate culture, or commitment to diversity that align with the candidate's profile.]

[Final paragraph: Reiterate interest in the position and express interest in an interview. Thank the employer for their time and consideration.]

Sincerely,


{candidate_name}
</output_format>

Generate the complete cover letter now, following the output format exactly and applying all requirements. Replace all bracketed instructions with actual content."""

    # System message for role-setting
    system_message = """You are an expert cover letter writer with 15 years of experience helping candidates land jobs at top companies across all industries. You excel at:

- Identifying key resume highlights that match job requirements
- Writing compelling narratives that showcase candidate strengths without exaggeration
- Adapting tone and style precisely to company culture and industry norms
- Maintaining appropriate length while maximizing impact and readability
- Using specific examples and concrete achievements rather than generic statements
- Crafting authentic, genuine language that sounds human and professional

You understand that cover letters should be concise, focused, and tailored to demonstrate clear value to the employer."""

    # Initialize client inside function to avoid initialization issues
    client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    message = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=1500,
        system=system_message,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    return message.content[0].text


def generate_statement_of_interest(resume_text, company_name, role_title, job_description=""):
    """Generate a brief 'why I want this job' statement using Claude Haiku."""

    prompt = f"""Based on the following information, write a brief 2-3 sentence statement explaining why the candidate wants this job. The statement should be honest, specific, and professional.

<resume>
{resume_text}
</resume>

<job_details>
Company: {company_name}
Role: {role_title}
Job Description: {job_description if job_description else "Not provided"}
</job_details>

Write a 2-3 sentence statement that:
- Highlights genuine interest based on the candidate's background
- Mentions specific aspects of the role or company that align with their experience
- Sounds authentic and not overly enthusiastic
- Can be refined into professional cover letter language

Output only the statement, no additional text or explanations."""

    client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    message = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=300,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    return message.content[0].text.strip()


def generate_application_answer(question, resume_text, company_name, role_title, job_description="", additional_context="", previous_responses="", question_notes="", resume_highlight=""):
    """Generate an answer to a random application question using Claude Haiku."""

    prompt = f"""You are helping a job candidate answer an application question. Based on the candidate's background and the job details, provide a professional, authentic answer.

<question>
{question}

CRITICAL: Your answer must DIRECTLY and EXPLICITLY answer this specific question above.
Do not answer a different question or go off-topic.
</question>"""

    # Add question-specific notes (context to incorporate)
    if question_notes:
        prompt += f"""

<candidate_notes_for_this_question>
The candidate provided these notes as CONTEXT and IDEAS to incorporate into your answer:

{question_notes}

Use these notes to enrich and add depth to your answer.
However, ensure you are ANSWERING THE QUESTION ABOVE, not just expanding on these notes.
</candidate_notes_for_this_question>"""

    prompt += f"""

<resume>
{resume_text}
</resume>"""

    # Add resume highlight if provided
    if resume_highlight:
        prompt += f"""

<resume_highlight>
The candidate specifically wants to EMPHASIZE these experiences/achievements from their resume:

{resume_highlight}

IMPORTANT: Make sure to feature and highlight these specific items in your answer when relevant to the question.
</resume_highlight>"""

    prompt += f"""

<job_details>
Company: {company_name}
Role: {role_title}
Job Description: {job_description if job_description else "Not provided"}
</job_details>

<additional_context>
{additional_context if additional_context else "No additional context provided."}
</additional_context>"""

    # Add previous responses if provided
    if previous_responses:
        prompt += f"""

<previous_responses>
You have already written the following for this application:

{previous_responses}

IMPORTANT: Avoid repeating the same experiences, skills, or examples mentioned above.
Highlight DIFFERENT aspects of the candidate's background.
Choose different stories, projects, or qualities to showcase.
Ensure this answer complements rather than duplicates what's already been written.
</previous_responses>"""

    prompt += f"""

REMINDER: You are answering this question: "{question}"

Write a clear, concise answer (2-4 sentences) that:
- DIRECTLY and EXPLICITLY answers the question above
- Uses specific examples from the candidate's background when relevant
- Incorporates the candidate's notes/context if provided
- Sounds authentic and professional
- Is appropriate for a job application
- Doesn't sound overly eager or generic

Output only the answer, no additional text or explanations."""

    client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    message = client.messages.create(
        model="claude-3-haiku-20240307",
        max_tokens=400,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    return message.content[0].text.strip()


def export_to_docx(cover_letter_text):
    """Export cover letter to .docx format with proper formatting."""
    doc = Document()

    # Set up document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Split cover letter into lines and add to document
    lines = cover_letter_text.split('\n')
    for line in lines:
        paragraph = doc.add_paragraph(line)
        # Set font size to 11pt
        for run in paragraph.runs:
            run.font.size = Pt(11)

    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_pdf(cover_letter_text):
    """Export cover letter to .pdf format with proper formatting."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Helvetica', '', 11)

    # Add all text at once using multi_cell
    pdf.multi_cell(0, 6, cover_letter_text)

    # Return PDF as bytes
    return bytes(pdf.output())


# ===== MAIN APP =====

# Show home page FIRST (before auth)
show_app = st.session_state.get("show_app", False)

if not show_app:
    st.title("AI-Powered Application Assistant")
    st.markdown("### Your intelligent companion for job applications")

    st.markdown("---")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("### 📝 Smart Cover Letters")
        st.markdown("""
        Generate professional, tailored cover letters in seconds.
        - Multiple length options
        - Customizable tone
        - Download as .txt, .docx, or .pdf
        """)

    with col2:
        st.markdown("### 💬 Answer Questions")
        st.markdown("""
        Craft compelling answers to application questions.
        - Context-aware responses
        - Avoid repetition across answers
        - Highlight specific experiences
        """)

    with col3:
        st.markdown("### 📊 Track & Save")
        st.markdown("""
        Keep your application materials organized.
        - Save multiple resumes
        - Cover letter history
        - Profile management
        """)

    st.markdown("---")

    st.markdown("### How It Works")
    st.markdown("""
    1. **Enter your details** - Add your resume, job info, and why you want the role
    2. **Generate instantly** - AI creates professional content tailored to the job
    3. **Download & apply** - Get your cover letter in any format and submit with confidence
    """)

    st.markdown("---")

    st.markdown("### Technical Implementation")
    st.markdown("""
    Built by **Nicolo Pastrone** with **Python** and **Streamlit** for a responsive web interface. Integrated **Anthropic's Claude API** for intelligent text generation with custom prompt engineering. **Supabase** handles authentication and PostgreSQL database with row-level security policies. Developed collaboratively with **Claude Code**, demonstrating modern AI-assisted development workflows.
    """)

    st.markdown("---")

    col_start1, col_start2, col_start3 = st.columns([1, 1, 1])

    with col_start2:
        if st.button("Get Started", type="primary", use_container_width=True):
            st.session_state["show_app"] = True
            st.rerun()

    st.caption("No account needed! Click 'Get Started' to try the app instantly.")

    st.stop()

# Check authentication (allow guest mode)
user_id = check_auth()
is_guest = st.session_state.get("guest_mode", False)

if not user_id and not is_guest:
    show_auth_page()
    st.stop()

st.title("AI-Powered Application Assistant")
st.caption("Your AI-powered job application toolkit")

# First-time user welcome banner
if st.session_state.get("first_time_user", True):
    col1, col2 = st.columns([5, 1])
    with col1:
        st.success("""
        **Welcome!** Here's how to get started:
        1. Enter your profile info in the sidebar (name, address, resume)
        2. Scroll down to fill in job details (company, role, job description)
        3. Click 'Generate Cover Letter' to create your tailored letter
        4. Download in your preferred format (.txt, .docx, or .pdf)

        Need help? Click 'How to Use This App' in the sidebar.
        """)
    with col2:
        if st.button("Got it"):
            st.session_state["first_time_user"] = False
            st.rerun()

# Show guest banner
if is_guest:
    st.info("You're using guest mode. Create an account to save your resumes, cover letters, and history!")

# Load profile (only for logged-in users)
profile = load_profile(user_id) if user_id else {
    "linkedin_url": "",
    "github_url": "",
    "portfolio_url": "",
    "candidate_name": "",
    "candidate_address": "",
    "default_length": "concise",
    "default_tone": "conversational"
}

# Sidebar for profile and resume management
with st.sidebar:
    # Home button
    if st.button("Home", use_container_width=True):
        st.session_state["show_app"] = False
        st.rerun()

    # Help section
    with st.expander("How to Use This App"):
        st.markdown("""
        **Quick Start:**
        1. Enter your name and address in 'Your Profile'
        2. Add your resume (paste or upload)
        3. Fill in job details below
        4. Click 'Generate Cover Letter'

        **Main Sections:**
        - **Job Details**: Company, role, job description
        - **Generate Cover Letter**: Create tailored cover letters
        - **Answer Application Question**: Handle essay questions

        **Tips for Best Results:**
        - Be specific in "Why do you want this job?"
        - Include the full job description when possible
        - Use "Resume Highlight" to emphasize specific experiences
        - Try different tones (Professional, Conversational, etc.)

        **Guest vs Account:**
        - Guests can generate and download everything
        - Accounts can save resumes and cover letter history
        """)

    st.divider()

    # Show login/logout based on mode
    if is_guest:
        st.info("Create an account to save your work!")
        if st.button("Sign Up / Login", use_container_width=True, type="primary"):
            if "guest_mode" in st.session_state:
                del st.session_state["guest_mode"]
            st.rerun()
    else:
        if st.button("Logout", use_container_width=True):
            logout_user()
            st.rerun()

    st.divider()

    # Section 1: Profile Links (only for logged-in users)
    if not is_guest:
        st.subheader("Profile Links")

        # Display saved links with copy functionality
        if profile.get("linkedin_url") or profile.get("github_url") or profile.get("portfolio_url"):
            if profile.get("linkedin_url"):
                st.markdown("**LinkedIn:**")
                st.code(profile['linkedin_url'], language=None)

            if profile.get("github_url"):
                st.markdown("**GitHub:**")
                st.code(profile['github_url'], language=None)

            if profile.get("portfolio_url"):
                st.markdown("**Portfolio:**")
                st.code(profile['portfolio_url'], language=None)
        else:
            st.info("No profile links saved yet.")

        # Edit profile links
        with st.expander("Edit Profile Links", expanded=False):
            linkedin_url = st.text_input("LinkedIn URL:", value=profile.get("linkedin_url", ""))
            github_url = st.text_input("GitHub URL:", value=profile.get("github_url", ""))
            portfolio_url = st.text_input("Portfolio URL:", value=profile.get("portfolio_url", ""))

            if st.button("Save Profile Links", use_container_width=True):
                profile["linkedin_url"] = linkedin_url
                profile["github_url"] = github_url
                profile["portfolio_url"] = portfolio_url
                save_profile(user_id, profile)
                st.success("Profile links saved!")
                st.rerun()

        st.divider()

    # Section 2: Resume Management (only for logged-in users)
    if not is_guest:
        st.subheader("Resume Management")

        # Load existing resumes
        saved_resumes = load_resumes(user_id)
    else:
        saved_resumes = []

    if saved_resumes:
        # Quick action: Use Latest Resume
        latest_resume = saved_resumes[0]  # First item (newest) since sorted desc
        st.info(f"Latest: {latest_resume['resume_name']}")
        if st.button("Use Latest Resume", use_container_width=True):
            st.session_state["resume_text"] = latest_resume["resume_text"]
            st.session_state["candidate_name"] = latest_resume["resume_name"]
            st.session_state["candidate_address"] = latest_resume.get("resume_address", "")
            st.success("Latest resume loaded!")
            st.rerun()

        st.divider()

        # Full resume selector
        resume_options = ["Enter new resume"] + [f"{r['resume_name']} - {r['date_saved']}" for r in saved_resumes]
        selected_resume = st.selectbox("Or select any resume:", resume_options)

        if selected_resume != "Enter new resume":
            resume_index = resume_options.index(selected_resume) - 1
            selected_resume_data = saved_resumes[resume_index]
            st.session_state["resume_text"] = selected_resume_data["resume_text"]
            st.session_state["candidate_name"] = selected_resume_data["resume_name"]
            st.session_state["candidate_address"] = selected_resume_data.get("resume_address", "")
    else:
        st.info("No saved resumes yet. Add your first resume below.")

    # Add/Update Resume section
    with st.expander("Add New Resume", expanded=False):
        upload_option = st.radio("How would you like to provide your resume?", ["Paste text", "Upload file"], horizontal=True)

        if upload_option == "Upload file":
            uploaded_file = st.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])
            if uploaded_file is not None:
                try:
                    if uploaded_file.name.endswith('.pdf'):
                        resume_text = extract_text_from_pdf(uploaded_file)
                    elif uploaded_file.name.endswith('.docx'):
                        resume_text = extract_text_from_docx(uploaded_file)
                    st.success("Resume uploaded successfully!")
                    st.session_state["resume_text"] = resume_text
                    st.session_state["uploaded_file"] = uploaded_file
                    st.session_state["uploaded_file_name"] = uploaded_file.name
                except Exception as e:
                    st.error(f"Error reading file: {str(e)}")
                    resume_text = ""
            else:
                resume_text = st.session_state.get("resume_text", "")
        else:
            resume_text = st.text_area(
                "Your Resume (paste full resume text):",
                value=st.session_state.get("resume_text", ""),
                height=150
            )
            # Clear uploaded file from session if pasting text
            if "uploaded_file" in st.session_state:
                del st.session_state["uploaded_file"]

        # Save resume button (only for logged-in users)
        if is_guest:
            st.info("Create an account to save resumes for later use.")
        else:
            if st.button("Save Resume", use_container_width=True):
                if not all([candidate_name, candidate_address, resume_text]):
                    st.error("Please fill in name, address, and resume text to save.")
                else:
                    resume_data = {
                        "resume_name": candidate_name,
                        "resume_address": candidate_address,
                        "resume_text": resume_text
                    }

                    save_resume(user_id, resume_data)
                    st.success(f"Resume saved! You now have {len(load_resumes(user_id))} saved resume(s).")

    if not is_guest:
        st.divider()

        # Section 4: Cover Letter History
        st.subheader("Cover Letter History")
        saved_cover_letters = load_cover_letters(user_id)
    else:
        saved_cover_letters = []
    if saved_cover_letters:
        st.caption(f"Total saved: {len(saved_cover_letters)}")
        for i, cl in enumerate(reversed(saved_cover_letters[-5:])):
            with st.expander(f"{cl['company']} - {cl['role']}", expanded=False):
                st.text(cl['cover_letter'][:200] + "...")
                st.caption(f"Created: {cl['date_created']}")

                # Download buttons in columns
                hist_col1, hist_col2, hist_col3 = st.columns(3)

                with hist_col1:
                    st.download_button(
                        label=".txt",
                        data=cl['cover_letter'],
                        file_name=f"cover_letter_{cl['company'].replace(' ', '_')}.txt",
                        mime="text/plain",
                        key=f"download_txt_{i}_{cl['date_created']}",
                        use_container_width=True
                    )

                with hist_col2:
                    docx_data = export_to_docx(cl['cover_letter'])
                    st.download_button(
                        label=".docx",
                        data=docx_data,
                        file_name=f"cover_letter_{cl['company'].replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_docx_{i}_{cl['date_created']}",
                        use_container_width=True
                    )

                with hist_col3:
                    pdf_data = export_to_pdf(cl['cover_letter'])
                    st.download_button(
                        label=".pdf",
                        data=pdf_data,
                        file_name=f"cover_letter_{cl['company'].replace(' ', '_')}.pdf",
                        mime="application/pdf",
                        key=f"download_pdf_{i}_{cl['date_created']}",
                        use_container_width=True
                    )
    else:
        st.info("No saved cover letters yet.")

# Main area - Job Details and Cover Letter Generation

# ===== SECTION 0: ENTER YOUR INFO =====
st.header("Step 1: Enter Your Info")
st.caption("Provide your profile information and resume")

# Create three columns for Name, Address, Resume
info_col1, info_col2, info_col3 = st.columns([1, 1, 2])

with info_col1:
    candidate_name = st.text_input(
        "Your Full Name:",
        value=st.session_state.get("candidate_name", ""),
        key="candidate_name_input",
        placeholder="e.g., Jane Smith"
    )
    st.session_state["candidate_name"] = candidate_name

with info_col2:
    candidate_address = st.text_area(
        "Your Address:",
        value=st.session_state.get("candidate_address", ""),
        height=80,
        help="Auto-filled from saved resumes but you can edit it.",
        key="candidate_address_input",
        placeholder="123 Main St\nBoston, MA 02101"
    )
    st.session_state["candidate_address"] = candidate_address

with info_col3:
    resume_text = st.text_area(
        "Your Resume:",
        value=st.session_state.get("resume_text", ""),
        height=80,
        help="Paste your resume text here, or use Resume Management in sidebar to save resumes for quick access.",
        key="resume_text_input",
        placeholder="Paste your resume text here..."
    )
    st.session_state["resume_text"] = resume_text

# Show status
if resume_text and candidate_name:
    st.success(f"Ready to generate cover letters for {candidate_name}")
elif not resume_text:
    st.info("Tip: If you have an account, use the Resume Management section in the sidebar to save and reuse resumes.")

st.divider()

# ===== SECTION 1: JOB DETAILS =====
st.header("Step 2: Enter Job Details")
st.caption("Provide information about the job you're applying for")

# Action buttons
action_col1, action_col2 = st.columns([1, 1])
with action_col1:
    if st.button("Use Sample Data", help="Fill in sample data to see how the app works"):
        st.session_state["candidate_name"] = "Jane Smith"
        st.session_state["candidate_address"] = "123 Main Street\nBoston, MA 02101"
        st.session_state["resume_text"] = """Jane Smith
123 Main Street, Boston, MA 02101
jane.smith@email.com | (555) 123-4567

EDUCATION
Bachelor of Science in Computer Science, Boston University - May 2023
GPA: 3.8/4.0

EXPERIENCE
Software Engineering Intern, TechCorp Inc. - Summer 2022
- Built a data pipeline that reduced processing time by 60%
- Collaborated with cross-functional teams to deliver features
- Wrote unit tests achieving 95% code coverage

Research Assistant, BU Computer Science Department - 2021-2023
- Developed machine learning models for natural language processing
- Published findings in undergraduate research symposium

SKILLS
Python, JavaScript, React, SQL, Git, AWS"""
        st.session_state["company_name_sample"] = "Anthropic"
        st.session_state["role_title_sample"] = "Software Engineer"
        st.session_state["job_description_sample"] = """We're looking for a software engineer to join our team working on AI safety and research. You'll be building tools that help make AI systems more helpful, harmless, and honest.

Requirements:
- Strong programming skills in Python
- Experience with modern web frameworks
- Passion for AI safety and alignment
- Excellent communication skills"""
        st.session_state["why_want_job_input"] = "I'm passionate about AI safety and want to contribute to building more reliable and beneficial AI systems. My experience in machine learning research and software development would allow me to make meaningful contributions to Anthropic's mission."
        st.success("Sample data loaded! Check the sidebar and form fields.")
        st.rerun()

with action_col2:
    if st.button("Clear Form", help="Clear all fields and start fresh"):
        # Clear all session state keys related to the form
        keys_to_clear = [
            "company_name_sample", "role_title_sample", "job_description_sample",
            "resume_text", "candidate_name", "candidate_address",
            "company_name_key", "role_title_key", "job_description_key",
            "additional_context_key", "resume_highlight_general", "why_want_job_input"
        ]
        for key in keys_to_clear:
            if key in st.session_state:
                del st.session_state[key]
        st.success("Form cleared!")
        st.rerun()

st.divider()

company_name = st.text_input(
    "Company Name:",
    value=st.session_state.get("company_name_sample", ""),
    key="company_name_key"
)
role_title = st.text_input(
    "Role/Position Title:",
    value=st.session_state.get("role_title_sample", ""),
    key="role_title_key"
)

job_description = st.text_area(
    "Job Description (optional but recommended):",
    value=st.session_state.get("job_description_sample", ""),
    height=150,
    help="Paste the job description here for better-tailored cover letters.",
    key="job_description_key"
)

additional_context = st.text_area(
    "Additional Context (optional):",
    height=100,
    help='Any extra context for the AI. e.g., "This is a general application, not for a specific role" or "I\'m currently working at X and looking to transition to Y"',
    placeholder='e.g., "This is a general application to the company, not a specific job posting."',
    key="additional_context_key"
)

resume_highlight = st.text_area(
    "Want to highlight anything specific from your resume? (optional):",
    height=80,
    placeholder='e.g., "My internship at X where I did Y" or "The data pipeline project that reduced processing time by 60%"',
    help="Specify particular experiences, projects, or achievements from your resume that you want emphasized in cover letters and answers.",
    key="resume_highlight_general"
)

st.divider()

# ===== SECTION 2: GENERATE COVER LETTER =====
st.header("Step 3: Generate Cover Letter")
st.caption("Configure preferences and generate a tailored cover letter")

# Preferences
st.subheader("Preferences")

col1, col2 = st.columns(2)

with col1:
    length_option = st.radio(
        "Letter Length:",
        ["Concise (200-325 words)", "Standard (325-450 words)"],
        index=0,
        help="Concise is recommended for most applications"
    )

with col2:
    tone_option = st.selectbox(
        "Tone:",
        [
            "Conversational - Warm but professional",
            "Professional - Formal and traditional",
            "Enthusiastic - Energetic and passionate",
            "Confident - Bold and direct"
        ],
        index=0,
        help="Match the tone to the company culture"
    )

st.subheader("Statement of Interest")

# Generate statement button (placed before text area so generated content shows up)
if st.button("Generate Statement", help="Auto-generate a 'why I want this job' statement based on your resume and the job details"):
    if not all([resume_text, company_name, role_title]):
        st.error("Please fill in your resume, company name, and role title to generate a statement.")
    else:
        with st.spinner("Generating statement..."):
            try:
                statement = generate_statement_of_interest(
                    resume_text,
                    company_name,
                    role_title,
                    job_description
                )
                st.session_state["why_want_job_input"] = statement
                st.success("Statement generated!")
                st.rerun()
            except Exception as e:
                st.error(f"Error generating statement: {str(e)}")

why_want_job = st.text_area(
    "Why do you want this job? (rough notes are fine):",
    value=st.session_state.get("why_want_job_input", ""),
    height=150,
    help="Be honest and specific. This will be refined into professional cover letter language.",
    key="why_want_job_input"
)

# Generate button
if st.button("Generate Cover Letter", type="primary"):
    if not all([candidate_name, candidate_address, resume_text, company_name, role_title, why_want_job]):
        st.error("Please fill in all required fields.")
    else:
        with st.spinner("Generating your cover letter..."):
            try:
                # Parse preferences
                length = "concise" if "Concise" in length_option else "standard"
                tone = tone_option.split(" - ")[0].lower()

                # Generate cover letter
                cover_letter = generate_cover_letter(
                    resume_text,
                    candidate_name,
                    candidate_address,
                    company_name,
                    role_title,
                    why_want_job,
                    job_description,
                    additional_context,
                    resume_highlight,
                    length,
                    tone
                )

                # Store in session state for rating
                st.session_state["last_cover_letter"] = cover_letter
                st.session_state["last_generation_data"] = {
                    "company": company_name,
                    "role": role_title,
                    "resume_text": resume_text,
                    "job_description": job_description,
                    "why_want_job": why_want_job
                }
                st.session_state["just_generated"] = True

                # Track for application session (to avoid repetition)
                if "application_session" not in st.session_state:
                    st.session_state["application_session"] = []
                st.session_state["application_session"].append({
                    "type": "cover_letter",
                    "content": cover_letter,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })

            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                st.info("Make sure your ANTHROPIC_API_KEY is set in the .env file.")

# Display cover letter if it exists in session state
if "last_cover_letter" in st.session_state and st.session_state["last_cover_letter"]:
    cover_letter = st.session_state["last_cover_letter"]
    gen_data = st.session_state.get("last_generation_data", {})

    # Show success message
    if st.session_state.get("just_generated", False):
        st.success("Cover letter generated!")
        st.session_state["just_generated"] = False
    elif st.session_state.get("cover_letter_saved", False):
        st.success("Cover letter saved!")
        st.session_state["cover_letter_saved"] = False

    # Display the cover letter
    st.subheader("Your Cover Letter")
    st.text_area("", value=cover_letter, height=500, key="generated_cl")

    # Download buttons
    st.subheader("Download")
    download_col1, download_col2, download_col3 = st.columns(3)

    with download_col1:
        st.download_button(
            label="Download as .txt",
            data=cover_letter,
            file_name=f"cover_letter_{gen_data.get('company', 'company').replace(' ', '_')}.txt",
            mime="text/plain",
            use_container_width=True
        )

    with download_col2:
        docx_data = export_to_docx(cover_letter)
        st.download_button(
            label="Download as .docx",
            data=docx_data,
            file_name=f"cover_letter_{gen_data.get('company', 'company').replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with download_col3:
        pdf_data = export_to_pdf(cover_letter)
        st.download_button(
            label="Download as .pdf",
            data=pdf_data,
            file_name=f"cover_letter_{gen_data.get('company', 'company').replace(' ', '_')}.pdf",
            mime="application/pdf",
            use_container_width=True
        )

    # Save button (only for logged-in users)
    st.divider()
    if is_guest:
        st.info("Create an account to save cover letters to your history!")
    else:
        if st.button("Save Cover Letter", use_container_width=True):
            cover_letter_data = {
                "company": gen_data.get("company", ""),
                "role": gen_data.get("role", ""),
                "cover_letter": cover_letter,
                "date_created": datetime.now().strftime("%Y-%m-%d %H:%M")
            }
            save_cover_letter(user_id, cover_letter_data)
            st.session_state["cover_letter_saved"] = True
            st.rerun()

    # Rating system (optional save for guests)
    if not is_guest:
        st.divider()
        st.subheader("Rate this cover letter")
        st.caption("Help improve future generations by rating this output")

        rating_col1, rating_col2 = st.columns(2)

        with rating_col1:
            if st.button("Good", use_container_width=True):
                rating_data = {
                    "rating": "good",
                    "cover_letter": cover_letter,
                    "resume_text": gen_data.get("resume_text", ""),
                    "job_description": gen_data.get("job_description", ""),
                    "why_want_job": gen_data.get("why_want_job", ""),
                    "company": gen_data.get("company", ""),
                    "role": gen_data.get("role", ""),
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                save_rating(user_id, rating_data)
                st.success("Thanks for your feedback!")

        with rating_col2:
            if st.button("Bad", use_container_width=True):
                rating_data = {
                    "rating": "bad",
                    "cover_letter": cover_letter,
                    "resume_text": gen_data.get("resume_text", ""),
                    "job_description": gen_data.get("job_description", ""),
                    "why_want_job": gen_data.get("why_want_job", ""),
                    "company": gen_data.get("company", ""),
                    "role": gen_data.get("role", ""),
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                save_rating(user_id, rating_data)
                st.info("Thanks for your feedback. We'll use this to improve!")

# Answer Application Question Section
st.divider()

# Header with Clear Session button
col1, col2 = st.columns([3, 1])
with col1:
    st.header("Step 4: Answer Application Question (Optional)")
    st.caption("Got a random question on the application? Generate a tailored answer based on your resume and job details.")
with col2:
    if st.button("Clear Session", help="Start fresh for a new application. Clears tracked responses to avoid repetition."):
        st.session_state["application_session"] = []
        st.success("Session cleared!")
        st.rerun()

# Show session info if exists
if "application_session" in st.session_state and st.session_state["application_session"]:
    with st.expander(f"Already written for this application ({len(st.session_state['application_session'])} item(s))", expanded=False):
        for idx, item in enumerate(st.session_state["application_session"]):
            st.markdown(f"**{idx + 1}. {item['type'].replace('_', ' ').title()}** ({item['timestamp']})")
            st.text(item['content'][:150] + "..." if len(item['content']) > 150 else item['content'])
            st.divider()

application_question = st.text_area(
    "Paste your application question here:",
    height=100,
    placeholder='e.g., "Why are you interested in working at [Company]?" or "Describe a time you overcame a challenge."',
    key="app_question_input"
)

question_notes = st.text_area(
    "Your notes / draft response (optional but recommended):",
    height=120,
    placeholder='e.g., "I\'ve been using this app for months and love it" or "I led a similar project at my previous company where..."',
    help="Add specific context, personal experiences, or rough notes related to this question. This helps the AI craft a more authentic, detailed answer.",
    key="question_notes_input"
)

# Checkbox to avoid repetition
avoid_repetition = st.checkbox(
    "Avoid repeating previous responses",
    value=True,
    help="When checked, the AI will avoid repeating experiences/skills from your cover letter or previous answers."
)

if st.button("Generate Answer", type="secondary"):
    if not application_question.strip():
        st.error("Please enter a question.")
    elif not all([resume_text, company_name, role_title]):
        st.error("Please make sure you have a resume loaded and company/role information filled in.")
    else:
        with st.spinner("Generating answer..."):
            try:
                # Build previous responses string if checkbox is checked
                previous_responses_text = ""
                if avoid_repetition and "application_session" in st.session_state and st.session_state["application_session"]:
                    previous_items = []
                    for item in st.session_state["application_session"]:
                        previous_items.append(f"--- {item['type'].replace('_', ' ').title()} ---\n{item['content']}\n")
                    previous_responses_text = "\n".join(previous_items)

                answer = generate_application_answer(
                    application_question,
                    resume_text,
                    company_name,
                    role_title,
                    job_description,
                    additional_context,
                    previous_responses_text,
                    question_notes,
                    resume_highlight
                )

                st.session_state["last_app_answer"] = answer
                st.session_state["last_app_question"] = application_question

                # Track this answer in application session
                if "application_session" not in st.session_state:
                    st.session_state["application_session"] = []
                st.session_state["application_session"].append({
                    "type": "application_question",
                    "content": f"Q: {application_question}\nA: {answer}",
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })

            except Exception as e:
                st.error(f"Error generating answer: {str(e)}")

# Display generated answer if it exists
if "last_app_answer" in st.session_state and st.session_state["last_app_answer"]:
    st.subheader("Generated Answer")
    st.text_area("", value=st.session_state["last_app_answer"], height=150, key="generated_answer_display")

    # Download button
    st.download_button(
        label="Download Answer",
        data=st.session_state["last_app_answer"],
        file_name=f"application_answer_{company_name.replace(' ', '_') if company_name else 'answer'}.txt",
        mime="text/plain"
    )
